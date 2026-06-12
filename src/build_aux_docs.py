"""
build_aux_docs.py -- create the EM&S submission companions:
  paper/Highlights.docx   (3-5 bullets, each <= 85 characters)
  paper/CoverLetter.docx  (addressed to the EM&S editor)

  python src/build_aux_docs.py
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PAPER = Path(r"C:/Users/vaspapa/Desktop/LakeForcing_OpenDrift/paper")
SERIF = "Times New Roman"
HLINK = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
URL_RE = re.compile(r"(https?://[^\s,)]+)")


def add_hyperlink(paragraph, url, text, size_hp="22"):
    """Append a clickable external hyperlink run (blue, underlined, TNR 11)."""
    r_id = paragraph.part.relate_to(url, HLINK, is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
    un = OxmlElement("w:u"); un.set(qn("w:val"), "single"); rpr.append(un)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), SERIF); rf.set(qn("w:hAnsi"), SERIF); rpr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), size_hp); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); paragraph._p.append(hl)

TITLE = ("LakeForcing-OpenDrift: an open, reproducible pipeline for generating "
         "hydrodynamic and wind-wave forcing of inland lakes to drive Lagrangian "
         "transport models")

HIGHLIGHTS = [
    "Open pipeline turns global open data into transport forcing for any inland lake",
    "Reusable sigma-layer to z-level coupling links Delft3D-FLOW/WAVE to OpenDrift",
    "Closed-lake Delft3D setup auto-built from HydroLAKES/GLOBathy/DAHITI and ERA5",
    "Exports CF NetCDF with 3-D currents, waves and surface Stokes drift",
    "Demonstrated unchanged on twelve lakes across all inhabited continents",
]


def set_base(doc):
    n = doc.styles["Normal"]
    n.font.name = SERIF; n.font.size = Pt(11)


def build_highlights():
    doc = Document(); set_base(doc)
    h = doc.add_paragraph(); r = h.add_run("Highlights"); r.bold = True
    r.font.size = Pt(13)
    sub = doc.add_paragraph()
    s = sub.add_run(TITLE); s.italic = True; s.font.size = Pt(10.5)
    doc.add_paragraph()
    for hl in HIGHLIGHTS:
        assert len(hl) <= 85, f"{len(hl)} > 85: {hl}"
        p = doc.add_paragraph(style="List Bullet"); p.add_run(hl)
    out = PAPER / "Highlights.docx"; doc.save(out)
    print("wrote", out.name, "| max chars:", max(len(h) for h in HIGHLIGHTS))


def build_cover_letter():
    doc = Document(); set_base(doc)
    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10, bold=False, italic=False):
        p = doc.add_paragraph(); p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        for k, part in enumerate(URL_RE.split(text)):   # turn URLs into hyperlinks
            if not part:
                continue
            if k % 2 == 1:
                trail = ""
                while part and part[-1] in ".,;:":
                    trail = part[-1] + trail; part = part[:-1]
                add_hyperlink(p, part, part)
                if trail:
                    p.add_run(trail)
            else:
                r = p.add_run(part); r.bold = bold; r.italic = italic
        return p

    para("12 June 2026", align=WD_ALIGN_PARAGRAPH.LEFT, after=6)
    para("To the Editors-in-Chief,\nEnvironmental Modelling & Software",
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

    para("Dear Editors,", align=WD_ALIGN_PARAGRAPH.LEFT, after=10)

    para(f"We are pleased to submit our manuscript, “{TITLE}”, for "
         "consideration in Environmental Modelling & Software as an original "
         "software publication.")

    para("Lagrangian particle tracking is a standard tool for studying the "
         "transport of floating material, but in inland lakes its use is blocked "
         "by the absence of ready-made forcing: global ocean reanalyses stop at "
         "the coast, and lake hydrodynamic models are built by hand, one waterbody "
         "at a time. This is a growing gap, because recent cross-national surveys "
         "show that lakes and reservoirs are among the most acutely plastic-polluted "
         "freshwater systems on Earth.")

    para("Our manuscript presents LakeForcing-OpenDrift, an open and reproducible "
         "Python pipeline that assembles bathymetry and meteorology from open "
         "global datasets, automatically configures and runs a coupled "
         "Delft3D-FLOW + Delft3D-WAVE (SWAN) simulation for an arbitrary lake, and "
         "exports CF-compliant NetCDF that drives the OpenDrift particle tracker "
         "without modification. The methodological core is a fully specified "
         "transform that bridges Delft3D's terrain-following sigma-layers to "
         "OpenDrift's fixed metric z-levels, together with the curvilinear-to-regular "
         "regridding, velocity rotation and surface Stokes-drift derivation needed "
         "to make lake hydrodynamics ingestible by a generic ocean tracker. We "
         "demonstrate the pipeline, unchanged, across twelve morphologically and "
         "climatically diverse lakes on all inhabited continents, from 36°S to 60°N.")

    para("We believe the work fits the scope of Environmental Modelling & Software: "
         "it is a reusable, openly released modelling tool with a clearly described "
         "method, demonstrated generality, and full software and data availability. "
         "It addresses a concrete and timely barrier to environmental transport "
         "modelling in freshwater systems. The generality is supported by a model-to-model "
         "benchmark against an expert-built reservoir model and by an independent comparison "
         "of the exported surface temperature against satellite observations for four further "
         "lakes. We note that the openly released twelve-lake forcing dataset could form the "
         "basis of a companion data descriptor (e.g. in Data in Brief), which we would be "
         "glad to prepare as a linked co-submission should the editors consider it "
         "appropriate.")

    para("This manuscript is original, has not been published previously, and is "
         "not under consideration for publication elsewhere. All authors have "
         "approved the manuscript and agree to its submission. The authors declare "
         "no competing interests. This research did not receive any specific grant "
         "from funding agencies in the public, commercial, or not-for-profit sectors; "
         "the work was carried out using the existing research infrastructure of "
         "CERTH-ITI. The source code is openly available at "
         "https://github.com/vaspapa79/LakeForcing-OpenDrift and archived on Zenodo "
         "(concept DOI: https://doi.org/10.5281/zenodo.20627160); the generated "
         "twelve-lake forcing dataset is distributed as release assets of the same archive.")

    para("Thank you for your consideration. We look forward to your response.",
         after=14)

    para("Sincerely,", align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    para("Vassilios Papaioannou, on behalf of all co-authors",
         align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    para("Information Technologies Institute, Centre for Research and Technology "
         "Hellas (CERTH-ITI)", align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    para("6th km Charilaou-Thermi, 57001 Thessaloniki, Greece",
         align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    para("vaspapa@iti.gr  |  Tel. +30 697 285 4287",
         align=WD_ALIGN_PARAGRAPH.LEFT, after=2)

    out = PAPER / "CoverLetter.docx"; doc.save(out)
    print("wrote", out.name)


if __name__ == "__main__":
    build_highlights()
    build_cover_letter()

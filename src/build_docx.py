"""
build_docx.py -- render paper/CAGEO_manuscript.md to a submission-ready .docx:
Times New Roman, justified body, formatted title + author block (superscripts),
headings, inline bold/italic/code, bullet lists, both pipe tables, numbered
display equations ([[EQ:name]] -> docs/eq/EQ_name.png) and the seven embedded,
captioned figures ([[FIG:name]]).

  python src/build_docx.py            ->  paper/CAGEO_manuscript_VP.docx
"""
import re
import sys
from pathlib import Path
from PIL import Image
from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

ROOT = Path(r"C:/Users/vaspapa/Desktop/LakeForcing_OpenDrift")
sys.path.insert(0, str(ROOT / "src"))
from omml_equations import oMath, EQ_NUM  # noqa: E402

MD = ROOT / "paper" / "CAGEO_manuscript.md"
DOCX = ROOT / "paper" / "CAGEO_manuscript_VP.docx"
SERIF = "Times New Roman"
CODE = RGBColor(0x33, 0x33, 0x33)

FIGS = {
    "architecture": (ROOT / "docs/figure_architecture.png", 1,
        "Architecture of the LakeForcing pipeline: single-purpose, "
        "file-coupled modules transform open global data into OpenDrift-ready "
        "CF-NetCDF; the σ-to-z coupling (cf_export.py) is the engine-to-tracker "
        "bridge."),
    "sigma_schematic": (ROOT / "docs/figure_sigma_schematic.png", 2,
        "The σ-layer to z-level coupling. (a) Delft3D-FLOW stores fields on "
        "terrain-following σ-layers whose depths vary with water level and "
        "bathymetry; (b) each σ-centre depth is reconstructed and the field "
        "interpolated to fixed metric z-levels, clamped at the surface (z = 0) "
        "and masked below the bed."),
    "vertical": (ROOT / "output/figure_vertical_sigma_z.png", 3,
        "Depth-resolved exported product for Lake Erken: surface-intensified "
        "warming and wind-driven current shear are preserved, and levels below "
        "the local bed are masked."),
    "forcing": (ROOT / "docs/figure_forcing_example.png", 4,
        "Exported surface forcing fields for a representative lake (Polyfytos): "
        "(a) mean current speed with vectors, (b) surface temperature, "
        "(c) significant wave height — all from a single coupled FLOW+WAVE run."),
    "map": (ROOT / "docs/figure_lake_map.png", 5,
        "Geographic distribution of the twelve demonstration lakes (circles, natural "
        "lakes; squares, reservoirs), spanning all inhabited continents from "
        "36°S to 60°N."),
    "demonstration": (ROOT / "output/figure_demonstration.png", 6,
        "The pipeline demonstrated across the twelve lakes: mean surface-current "
        "speed (per-lake colour scale) with 36 h particle trajectories (black), "
        "release point (star) and endpoints (dots); land is grey."),
    "scatter": (ROOT / "docs/figure_drift_scatter.png", 7,
        "Physical consistency: 36 h mean drift versus (a) mean depth and "
        "(b) maximum surface current for the twelve lakes, read as rank tendencies "
        "(n = 12; Section 5.3)."),
    "validation": (ROOT / "docs/figure_validation.png", 8,
        "Consistency check against the expert-built Polyfytos model on the shared "
        "grid (48 h mean surface fields): (a, b) surface-current speed and "
        "(c) scatter; (d, e) surface temperature and (f) scatter (Section 5.6)."),
    "satellite": (ROOT / "docs/figure_satellite.png", 9,
        "Validation against satellite lake surface water temperature (Landsat-8/9 "
        "Collection-2 Level-2 thermal band, overpasses around 1-3 July 2022): per "
        "lake, satellite skin temperature, exported model surface temperature at "
        "the diurnal peak, and their scatter (Section 5.7)."),
}


def img_width_in(path, cap):
    im = Image.open(path)
    dpi = im.info.get("dpi", (96, 96))[0] or 96
    return min(im.size[0] / dpi, cap)


def spacer(doc, pt=3):
    """A minimal blank line (much smaller than an empty default paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.add_run("").font.size = Pt(pt)
    return p


URL_RE = re.compile(r"(https?://[^\s)]+)")


def _emit_plain(par, s):
    """Add plain text, turning any URL into a live (blue, underlined) hyperlink."""
    for k, part in enumerate(URL_RE.split(s)):
        if not part:
            continue
        if k % 2 == 1:                       # captured URL
            trail = ""
            while part and part[-1] in ".,;:":
                trail = part[-1] + trail; part = part[:-1]
            add_hyperlink(par, part, part)
            if trail:
                par.add_run(trail)
        else:
            par.add_run(part)


def add_runs(par, text):
    par.paragraph_format.space_after = Pt(3)

    def clean(x):
        return x.replace("\x00", "*").replace("\x01", "|")
    text = text.replace(r"\*", "\x00").replace(r"\|", "\x01")
    for tok in re.split(r"(\*\*.+?\*\*|`[^`]+?`|_\{[^}]*\}|\*[^*]+?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(clean(tok[2:-2])); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(clean(tok[1:-1])); r.font.name = SERIF
            r.font.size = Pt(10); r.font.color.rgb = RGBColor(0, 0, 0)
        elif tok.startswith("_{") and tok.endswith("}"):
            r = par.add_run(clean(tok[2:-1])); r.font.subscript = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(clean(tok[1:-1])); r.italic = True
        else:
            _emit_plain(par, clean(tok))


def front_matter(doc):
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LakeForcing: a σ-to-z coupling algorithm and open "
                  "pipeline for hydrodynamic and wind-wave forcing of inland lakes "
                  "to drive Lagrangian transport models")
    r.bold = True; r.font.size = Pt(16)

    authors = [("Vassilios Papaioannou", "1,*"), ("Christos G. E. Anagnostopoulos", "1"),
               ("Dimitrios Valsamis", "1"), ("Anastasia Moumtzidou", "1"),
               ("Ilias Gialampoukidis", "1"), ("Stefanos Vrochidis", "1"),
               ("Ioannis Kompatsiaris", "1")]
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (name, sup) in enumerate(authors):
        if i:
            p.add_run(", ").font.size = Pt(12)
        p.add_run(name).font.size = Pt(12)
        s = p.add_run(sup); s.font.superscript = True; s.font.size = Pt(12)

    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    su = a.add_run("1"); su.font.superscript = True
    ar = a.add_run(" Information Technologies Institute, Centre for Research and "
                   "Technology Hellas (CERTH-ITI), 6th km Charilaou-Thermi, "
                   "57001 Thessaloniki, Greece")
    ar.italic = True; ar.font.size = Pt(10)

    orcids = [("Vassilios Papaioannou", "0000-0001-6598-7107"),
              ("Christos G. E. Anagnostopoulos", "0009-0006-9839-4870"),
              ("Dimitrios Valsamis", "0009-0002-6699-4426"),
              ("Anastasia Moumtzidou", "0000-0001-7615-8400"),
              ("Ilias Gialampoukidis", "0000-0002-5234-9795"),
              ("Stefanos Vrochidis", "0000-0002-2505-9178"),
              ("Ioannis Kompatsiaris", "0000-0001-6447-9020")]
    o = doc.add_paragraph(); o.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lbl = o.add_run("ORCID iDs: "); lbl.italic = True; lbl.font.size = Pt(10)
    for k, (name, oid) in enumerate(orcids):
        if k:
            o.add_run("; ").font.size = Pt(10)
        o.add_run(f"{name} ").font.size = Pt(10)
        add_hyperlink(o, f"https://orcid.org/{oid}", oid)

    emails = [("Vassilios Papaioannou", "vaspapa@iti.gr"),
              ("Christos G. E. Anagnostopoulos", "anagn_c@iti.gr"),
              ("Dimitrios Valsamis", "dvalsamis@iti.gr"),
              ("Anastasia Moumtzidou", "moumtzid@iti.gr"),
              ("Ilias Gialampoukidis", "heliasgj@iti.gr"),
              ("Stefanos Vrochidis", "stefanos@iti.gr"),
              ("Ioannis Kompatsiaris", "ikom@iti.gr")]
    e = doc.add_paragraph(); e.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lbl = e.add_run("E-mail addresses: "); lbl.italic = True; lbl.font.size = Pt(10)
    for i, (name, mail) in enumerate(emails):
        if i:
            e.add_run("; ").font.size = Pt(10)
        e.add_run(f"{mail} ({name})").font.size = Pt(10)

    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run("* Corresponding author: Vassilios Papaioannou, vaspapa@iti.gr, "
              "Tel. +30 697 285 4287").font.size = Pt(10)
    doc.add_paragraph()


def emit_equation(doc, name):
    """Insert a native Word (OMML) equation, centred, with a right-flush number."""
    num = EQ_NUM.get(name, 0)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = (Inches(0.25), Inches(6.5), Inches(0.45))
    for col, w in zip(tbl.columns, widths):
        col.width = w
    for row in tbl.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
    mid = tbl.rows[0].cells[1]
    mid.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pm = mid.paragraphs[0]; pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pm._p.append(parse_xml(oMath(name)))          # editable Word equation
    right = tbl.rows[0].cells[2]
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pr = right.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pr.add_run(f"({num})"); rr.font.name = SERIF
    spacer(doc)


def emit_figure(doc, name):
    path, num, caption = FIGS[name]
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(img_width_in(path, 6.9)))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2); cap.paragraph_format.space_before = Pt(2)
    b = cap.add_run(f"Figure {num}. "); b.bold = True; b.font.size = Pt(10)
    cr = cap.add_run(caption); cr.font.size = Pt(10)
    spacer(doc)


def emit_graphical_abstract(doc):
    """Insert the plate-assembled graphical abstract, full width, unnumbered."""
    path = ROOT / "docs/graphical_abstract.png"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(img_width_in(path, 6.9)))
    spacer(doc)


def _shade(cell, hexfill):
    """Set a table cell background fill (e.g. 'D9D9D9')."""
    tcpr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hexfill}"/>')
    tcpr.append(shd)


def _fix_widths(tbl, widths):
    """Pin column widths reliably (set on both columns and every cell)."""
    tbl.autofit = False
    for col, w in zip(tbl.columns, widths):
        col.width = w
    for row in tbl.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def emit_table(doc, rows):
    header = [c.strip() for c in rows[0]]
    # Header-less 2-column key/value table (e.g. Code metadata): no blank header
    # row; narrow shaded label column + wide value column.
    if len(header) == 2 and all(h == "" for h in header):
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r in rows[2:]:
            cells = tbl.add_row().cells
            lp = cells[0].paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(lp, r[0].strip())
            for rr in lp.runs:
                rr.bold = True; rr.font.color.rgb = RGBColor(0, 0, 0)
            _shade(cells[0], "F2F2F2")
            vp = cells[1].paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(vp, r[1].strip() if len(r) > 1 else "")
        _fix_widths(tbl, (Inches(2.15), Inches(5.05)))
        spacer(doc)
        return
    aligns = ["r" if c.strip().endswith(":") and not c.strip().startswith(":")
              else "l" for c in rows[1]]
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = "Table Grid"                       # black borders, no accent-blue
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]; cell.paragraphs[0].text = ""
        _shade(cell, "D9D9D9")                     # light-grey header, black text
        add_runs(cell.paragraphs[0], h)
        for rr in cell.paragraphs[0].runs:
            rr.bold = True; rr.font.color.rgb = RGBColor(0, 0, 0)
    for r in rows[2:]:
        cells = tbl.add_row().cells
        for j in range(len(header)):
            par = cells[j].paragraphs[0]
            add_runs(par, r[j].strip() if j < len(r) else "")
            par.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if aligns[j] == "r"
                             else WD_ALIGN_PARAGRAPH.LEFT)
    spacer(doc)


HLINK = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"


def add_hyperlink(paragraph, url, text):
    """Append a clickable external hyperlink run (blue, underlined, TNR 10)."""
    r_id = paragraph.part.relate_to(url, HLINK, is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
    un = OxmlElement("w:u"); un.set(qn("w:val"), "single"); rpr.append(un)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), SERIF); rf.set(qn("w:hAnsi"), SERIF); rpr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20"); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text
    t.set(qn("xml:space"), "preserve"); run.append(t)
    hl.append(run); paragraph._p.append(hl)


def emit_reference(doc, text):
    """One reference: justified, hanging indent; DOI URL as a live hyperlink."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3); pf.first_line_indent = Inches(-0.3)
    pf.space_after = Pt(4)
    m = re.search(r"(https://doi\.org/\S+)", text)
    if m:
        add_runs(p, text[:m.start()])
        add_hyperlink(p, m.group(1).rstrip("."), m.group(1).rstrip("."))
        if text[m.end():].strip():
            add_runs(p, text[m.end():])
    else:
        add_runs(p, text)


def split_row(line):
    line = line.strip().strip("|").replace(r"\|", "\x01")
    return [c.replace("\x01", r"\|") for c in line.split("|")]


def emit_article_info(doc, keywords, credit, abstract):
    """C&G front matter: ARTICLE INFO (Keywords, Authorship statement) + ABSTRACT."""
    doc.add_heading("ARTICLE INFO", level=1)
    kp = doc.add_paragraph(); kp.add_run("Keywords:").bold = True   # C&G: keywords as a list
    for kw in [k.strip() for k in keywords.split(";") if k.strip()]:
        doc.add_paragraph(kw)
    doc.add_heading("Authorship contribution statement", level=2)
    if credit:
        add_runs(doc.add_paragraph(), credit)
    doc.add_heading("ABSTRACT", level=1)
    add_runs(doc.add_paragraph(), abstract)


def main():
    lines = MD.read_text(encoding="utf-8").splitlines()
    start = next(i for i, l in enumerate(lines) if l.strip() == "## Abstract")
    body = lines[start:]

    doc = Document()
    # Computers & Geosciences Word template: single column, 1-inch margins,
    # Times New Roman, double-spaced, continuous line numbers for review.
    for sec in doc.sections:
        sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)
        sectPr = sec._sectPr
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1"); ln.set(qn("w:restart"), "continuous")
        ln.set(qn("w:distance"), "360")
        cols = sectPr.find(qn("w:cols"))
        (cols.addprevious(ln) if cols is not None else sectPr.append(ln))
    # base style: Times New Roman, justified, DOUBLE-spaced (C&G requirement)
    normal = doc.styles["Normal"]
    normal.font.name = SERIF; normal.font.size = Pt(10)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 2.0; pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    hsizes = {"Heading 1": 12, "Heading 2": 11, "Heading 3": 11}
    for hs in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            st = doc.styles[hs]
            st.font.name = SERIF
            st.font.color.rgb = RGBColor(0, 0, 0)   # black headings, no accent-blue
            st.font.bold = True
            if hs in hsizes:
                st.font.size = Pt(hsizes[hs])
            st.paragraph_format.line_spacing = 2.0
            st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            st.paragraph_format.space_before = Pt(6)
            st.paragraph_format.space_after = Pt(0)
        except KeyError:
            pass
    # Pre-parse the CRediT statement so it can be emitted in the ARTICLE INFO
    # block (template order) and skipped where it appears later in the body.
    credit_text = ""
    for k, l in enumerate(body):
        if l.strip().startswith("## CRediT"):
            j = k + 1; buf = []
            while j < len(body) and not body[j].strip().startswith("## "):
                if body[j].strip():
                    buf.append(body[j].strip())
                j += 1
            credit_text = " ".join(buf)
            break
    front_matter(doc)

    i = 0
    in_refs = False
    while i < len(body):
        s = body[i].strip()
        if not s or s == "---":
            i += 1; continue
        if s == "## Abstract":
            j = i + 1; abs_buf = []
            while j < len(body) and not body[j].strip().startswith("**Keywords"):
                if body[j].strip() and body[j].strip() != "---":
                    abs_buf.append(body[j].strip())
                j += 1
            kw_buf = []                                    # keywords may wrap over lines
            if j < len(body):
                kw_buf.append(re.sub(r"^\*\*Keywords:\*\*\s*", "", body[j].strip()))
                j += 1
                while j < len(body):
                    nx = body[j].strip()
                    if not nx or nx == "---" or nx.startswith(("#", "[[", "|", "- ")):
                        break
                    kw_buf.append(nx); j += 1
            emit_article_info(doc, " ".join(kw_buf), credit_text, " ".join(abs_buf))
            while i < len(body) and not re.match(r"## 1\.", body[i].strip()):
                i += 1
            continue
        if s.startswith("## CRediT"):                 # moved into ARTICLE INFO
            i += 1
            while i < len(body) and not body[i].strip().startswith("## "):
                i += 1
            continue
        m = re.match(r"\[\[EQ:([a-z]+)\]\]", s)
        if m:
            emit_equation(doc, m.group(1)); i += 1; continue
        m = re.match(r"\[\[FIG:([a-z_]+)\]\]", s)
        if m:
            emit_figure(doc, m.group(1)); i += 1; continue
        if s == "[[GRAPHICAL_ABSTRACT]]":
            emit_graphical_abstract(doc); i += 1; continue
        if s.startswith("|") and i + 1 < len(body) and set(body[i + 1].strip()) <= set("|-: "):
            rows = []
            while i < len(body) and body[i].strip().startswith("|"):
                rows.append(split_row(body[i])); i += 1
            emit_table(doc, rows); continue
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=3); i += 1; continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=2); i += 1; continue
        if s.startswith("## "):
            in_refs = (s[3:].strip().lower() == "references")
            doc.add_heading(s[3:], level=1); i += 1; continue
        if in_refs:
            emit_reference(doc, s); i += 1; continue
        if s.startswith("- "):
            buf = [s[2:]]; i += 1                       # gather wrapped continuation lines
            while i < len(body):
                nx = body[i].strip()
                if not nx or nx == "---" or nx.startswith(("#", "- ", "|", "[[")):
                    break
                buf.append(nx); i += 1
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, " ".join(buf)); continue
        # paragraph (gather wrapped lines)
        buf = [s]; i += 1
        while i < len(body):
            nx = body[i].strip()
            if not nx or nx == "---" or nx.startswith(("#", "- ", "|", "[[")):
                break
            buf.append(nx); i += 1
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, " ".join(buf))

    target = DOCX
    try:
        doc.save(target)
    except PermissionError:
        target = DOCX.with_name("CAGEO_manuscript_VP_uplift.docx")
        doc.save(target)
        print(f"NOTE: {DOCX.name} was locked (open in Word) -> saved to {target.name}")
    nf = sum(1 for l in body if l.strip().startswith("[[FIG:"))
    ne = sum(1 for l in body if l.strip().startswith("[[EQ:"))
    print(f"wrote {target} | {ne} equations, {nf} figures, {len(doc.tables)} tables")


if __name__ == "__main__":
    main()
